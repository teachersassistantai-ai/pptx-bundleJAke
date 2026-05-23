# =============================================================================
# Literacy & Numeracy Work Pack -- v4 RENDERER (GitHub-hosted, fetched + exec'd
# by the tiny bootstrapper step inside the .rai).
#
# The bootstrapper provides these names in the exec namespace:
#   lit_raw, num_raw, intro_raw, enrich_raw, seed_raw  (strings from {{...}})
#   topic, yl, lit_level, ability, timeframe           (param strings)
#   openai_key                                         (secrets.openai)
#   insert_temp_file                                   (Relevance builtin)
# This file sets RESULT to the temp-file URL (or an ERR string).
#
# Image is generated IN-PROCESS (OpenAI -> base64 -> BytesIO -> add_picture):
# no temp-URL round trip, which is what was silently dropping the cover.
# =============================================================================
import json, re, math, io, base64, warnings, requests, json_repair
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as WA
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml

warnings.filterwarnings("ignore", message="You have opted to not authenticate on client instantiation")


# ---------- robust JSON parse -----------------------------------------------
# LLMs (esp. via Relevance's Anthropic path, which doesn't enforce json_object)
# wrap output in ```json fences AND put literal newlines/tabs inside string
# values -- both invalid JSON. This parser strips fences and escapes literal
# control chars inside strings before parsing, plus a brace-extract fallback.
def _strip_fences(t):
    t = t.strip()
    if t.startswith("```"):
        t = re.sub(r"^```[a-zA-Z0-9]*[ \t]*\n?", "", t)
        t = re.sub(r"\n?[ \t]*```\s*$", "", t)
    return t.strip()


def _esc_ctrl(t):
    # escape raw newline/tab/CR that appear INSIDE a JSON string literal
    out = []
    in_str = False
    esc = False
    for ch in t:
        if esc:
            out.append(ch); esc = False; continue
        if ch == "\\":
            out.append(ch); esc = True; continue
        if ch == '"':
            in_str = not in_str; out.append(ch); continue
        if in_str and ch == "\n":
            out.append("\\n"); continue
        if in_str and ch == "\t":
            out.append("\\t"); continue
        if in_str and ch == "\r":
            out.append("\\r"); continue
        out.append(ch)
    return "".join(out)


def pj(raw):
    if not raw:
        return None
    if isinstance(raw, dict):
        return pj(raw["answer"]) if set(raw.keys()) == {"answer"} else raw
    if isinstance(raw, list) and raw and isinstance(raw[0], dict):
        return raw[0]
    t = str(raw).strip()
    if t.lower() in ("null", "none", ""):
        return None
    # json_repair fixes fences, unescaped quotes (dialogue!), literal control chars,
    # trailing commas -- everything LLMs get wrong. Try strict first (fast path).
    r = None
    try:
        r = json.loads(_strip_fences(t))
    except Exception:
        try:
            r = json_repair.loads(t)
        except Exception:
            r = None
    if isinstance(r, dict):
        return pj(r) if set(r.keys()) == {"answer"} else (r or None)
    if isinstance(r, list) and r and isinstance(r[0], dict):
        return r[0]
    return None


# ---------- year band + palette ----------------------------------------------
BANDS = {
    "PRIMARY": {"bf": "Bookman Old Style", "hf": "Bookman Old Style", "bs": 14, "h0": 34,
                "pri": (0xE5, 0x6B, 0x6F), "sec": (0xF4, 0xA2, 0x61), "ink": (0x33, 0x2A, 0x2A),
                "tint": "FDEBED", "paper": "FFFBF5", "line_h": 26, "box": "chunky", "cov_w": 12.0},
    "MIDDLE":  {"bf": "Calibri", "hf": "Calibri", "bs": 12, "h0": 28,
                "pri": (0x2E, 0x6F, 0xB5), "sec": (0x16, 0xA0, 0x85), "ink": (0x2D, 0x37, 0x48),
                "tint": "EAF2FB", "paper": "FFFFFF", "line_h": 18, "box": "clean", "cov_w": 11.0},
    "SENIOR":  {"bf": "Georgia", "hf": "Calibri", "bs": 11, "h0": 23,
                "pri": (0x1F, 0x2D, 0x3D), "sec": (0x9A, 0x6B, 0x3A), "ink": (0x22, 0x26, 0x2B),
                "tint": "F2F4F6", "paper": "FFFFFF", "line_h": 14, "box": "minimal", "cov_w": 9.0},
}

digits = "".join(c for c in yl if c.isdigit())
yr_num = int(digits[:2]) if digits else 7
band_name = "PRIMARY" if yr_num <= 4 else "MIDDLE" if yr_num <= 8 else "SENIOR"
B = BANDS[band_name]
PRI = RGBColor(*B["pri"]); SEC = RGBColor(*B["sec"]); INK = RGBColor(*B["ink"])
GY = RGBColor(0x6B, 0x70, 0x76); WHITE = RGBColor(0xFF, 0xFF, 0xFF)
PRI_HEX = "%02X%02X%02X" % B["pri"]; SEC_HEX = "%02X%02X%02X" % B["sec"]
TIGHT = band_name == "SENIOR"
GRID = "D9DCE1"  # soft grey gridlines (replaces default black)

# parse inputs
lit = pj(lit_raw); num = pj(num_raw)
enrich = pj(enrich_raw) or {}
topic_clarified = (enrich.get("topic_clarified") or topic or "").strip() or topic
au_context = (enrich.get("australian_context") or "").strip()
# Catchy generated booklet title (NOT the raw user topic, which may be dull/misspelt).
title = (enrich.get("pack_title") or "").strip() or topic
intro = (intro_raw or "").strip()
_tf = (timeframe or "").lower()
# True multi-day only: "1 day"/"2 days"/"3 days"/weeks. "Half day" is single-session.
is_multi = ("week" in _tf) or ("day" in _tf and "half" not in _tf)


def _ok(d):
    return isinstance(d, dict) and d.get("activities")


if not _ok(lit) and not _ok(num):
    RESULT = "ERR parse. Lit(%d): %s | Num(%d): %s" % (
        len(str(lit_raw)), str(lit_raw)[:300], len(str(num_raw)), str(num_raw)[:200])
else:
    doc = Document()

    # ---- base styles ----
    sn = doc.styles["Normal"]
    sn.font.name = B["bf"]; sn.font.size = Pt(B["bs"]); sn.font.color.rgb = INK
    sn.paragraph_format.line_spacing = 1.12
    sn.paragraph_format.space_after = Pt(4)
    for i, sb, sa in ((1, 18, 6), (2, 14, 4), (3, 10, 3)):
        s = doc.styles["Heading %d" % i]
        s.font.name = B["hf"]; s.font.color.rgb = PRI
        s.paragraph_format.space_before = Pt(sb); s.paragraph_format.space_after = Pt(sa)
    for sec in doc.sections:
        sec.top_margin = Cm(1.9); sec.bottom_margin = Cm(1.9)
        sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
    doc.sections[0].different_first_page_header_footer = True

    # ---------- low-level helpers ----------
    def shd(cell, hexclr):
        cell._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:val="clear" w:fill="%s"/>' % (nsdecls("w"), hexclr)))

    def cell_pad(cell, t=90, b=90, l=140, r=140):
        cell._tc.get_or_add_tcPr().append(parse_xml(
            '<w:tcMar %s><w:top w:w="%d" w:type="dxa"/><w:bottom w:w="%d" w:type="dxa"/>'
            '<w:left w:w="%d" w:type="dxa"/><w:right w:w="%d" w:type="dxa"/></w:tcMar>'
            % (nsdecls("w"), t, b, l, r)))

    def cell_borders(cell, left=None, box=None, sz=4):
        parts = []
        if left:
            parts.append('<w:left w:val="single" w:sz="26" w:space="0" w:color="%s"/>' % left)
        if box:
            for edge in ("top", "bottom", "right"):
                parts.append('<w:%s w:val="single" w:sz="%d" w:space="0" w:color="%s"/>' % (edge, sz, box))
            if not left:
                parts.append('<w:left w:val="single" w:sz="%d" w:space="0" w:color="%s"/>' % (sz, box))
        cell._tc.get_or_add_tcPr().append(parse_xml('<w:tcBorders %s>%s</w:tcBorders>' % (nsdecls("w"), "".join(parts))))

    def table_grid(tbl, color=GRID, sz=4):
        tblPr = tbl._tbl.tblPr
        tblPr.append(parse_xml(
            '<w:tblBorders %s>'
            '<w:top w:val="single" w:sz="%d" w:color="%s"/><w:left w:val="single" w:sz="%d" w:color="%s"/>'
            '<w:bottom w:val="single" w:sz="%d" w:color="%s"/><w:right w:val="single" w:sz="%d" w:color="%s"/>'
            '<w:insideH w:val="single" w:sz="%d" w:color="%s"/><w:insideV w:val="single" w:sz="%d" w:color="%s"/>'
            '</w:tblBorders>' % (nsdecls("w"), sz, color, sz, color, sz, color, sz, color, sz, color, sz, color)))
        cantsplit(tbl)

    _RX = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*)")

    def emit(p, text, sz=None, base_italic=False, clr=None, bold_all=False):
        sz = sz or B["bs"]
        for tok in _RX.split(str(text)):
            if not tok:
                continue
            if tok.startswith("**") and tok.endswith("**"):
                r = p.add_run(tok[2:-2]); r.bold = True
            elif tok.startswith("*") and tok.endswith("*"):
                r = p.add_run(tok[1:-1]); r.italic = True
            else:
                r = p.add_run(tok.replace("**", "")); r.italic = base_italic; r.bold = bold_all
            r.font.size = Pt(sz); r.font.name = B["bf"]
            if clr:
                r.font.color.rgb = clr
        return p

    def pa(text, bold=False, sz=None, it=False, clr=None, al=None, after=None, before=None):
        p = doc.add_paragraph()
        emit(p, text, sz=sz or B["bs"], base_italic=it, clr=clr, bold_all=bold)
        if al is not None:
            p.alignment = al
        if after is not None:
            p.paragraph_format.space_after = Pt(after)
        if before is not None:
            p.paragraph_format.space_before = Pt(before)
        return p

    def addh(text, lvl):
        h = doc.add_heading("", level=lvl)
        emit(h, text, sz=None, clr=PRI)
        h.paragraph_format.keep_with_next = True
        return h

    def cantsplit(tbl):
        for row in tbl.rows:
            row._tr.get_or_add_trPr().append(parse_xml("<w:cantSplit %s/>" % nsdecls("w")))

    def kicker(text):
        p = doc.add_paragraph()
        r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(B["bs"] - 4)
        r.font.color.rgb = SEC; r.font.name = B["hf"]
        r._r.get_or_add_rPr().append(parse_xml('<w:spacing %s w:val="40"/>' % nsdecls("w")))
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        return p

    def chip(p, text, fill):
        minimal = B["box"] == "minimal"
        r = p.add_run((" %s " % text.upper()) if not minimal else (text.upper() + "  "))
        r.bold = True; r.font.size = Pt(B["bs"] - 3); r.font.name = B["hf"]
        rPr = r._r.get_or_add_rPr()
        rPr.append(parse_xml('<w:spacing %s w:val="30"/>' % nsdecls("w")))
        if minimal:
            r.font.color.rgb = RGBColor(*B["sec"])
        else:
            r.font.color.rgb = WHITE
            rPr.append(parse_xml('<w:shd %s w:val="clear" w:fill="%s"/>' % (nsdecls("w"), fill)))

    # box purpose registry: (fill, accent, label, italic)
    BOXES = {
        "passage":    (B["tint"], PRI_HEX, None, False),
        "wordbank":   ("FFFFFF", SEC_HEX, "WORD BANK", False),
        "glossary":   ("FFFFFF", SEC_HEX, "WORDS TO KNOW", False),
        "fact":       ("FFF8E7", "E8A33D", "DID YOU KNOW?", True),
        "challenge":  ("FFF1E6", "E2703A", "CHALLENGE", False),
        "reflection": (B["tint"], PRI_HEX, "THINK ABOUT IT", True),
        "task":       (B["tint"], PRI_HEX, "YOUR TASK", False),
        "plain":      (B["tint"], PRI_HEX, None, False),
    }

    def box(text, kind, rows=None):
        fill, accent, label, ital = BOXES[kind]
        if B["box"] == "minimal":
            fill = "FFFFFF"  # senior: rule-only, no fill
        t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.LEFT
        cl = t.rows[0].cells[0]; cl.text = ""
        p = cl.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        if label:
            chip(p, label, accent)
            p.add_run().add_break()
        if rows:
            for j, line in enumerate(rows):
                pp = p if (not label and j == 0) else cl.add_paragraph()
                emit(pp, line, sz=B["bs"] - 1, base_italic=ital)
        else:
            emit(p, text, sz=B["bs"] - 1, base_italic=ital)
        if B["box"] != "minimal":
            shd(cl, fill)
        cell_borders(cl, left=accent, box=GRID)
        cell_pad(cl)
        if kind != "passage":  # short boxes shouldn't split across a page; let long passages flow
            cantsplit(t)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def writelines(n=2):
        n = max(1, (n - 1) if TIGHT else n)
        for _ in range(n):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(B["line_h"]); p.paragraph_format.space_after = Pt(0)
            p._p.get_or_add_pPr().append(parse_xml(
                '<w:pBdr %s><w:bottom w:val="dotted" w:sz="6" w:space="2" w:color="B5B9BF"/></w:pBdr>' % nsdecls("w")))

    def gq(q):
        return str(q.get("q", q.get("question", ""))) if isinstance(q, dict) else str(q)

    def gl(q):
        return q.get("lines", 2) if isinstance(q, dict) else 2

    def numq(i, text):
        p = doc.add_paragraph(); p.paragraph_format.keep_with_next = True
        r = p.add_run("%d.  " % i); r.bold = True; r.font.color.rgb = PRI; r.font.name = B["hf"]
        emit(p, text)
        return p

    def hdr_cells(cells, labels):
        for ci, lbl in enumerate(labels):
            cl = cells[ci]; cl.text = ""
            r = cl.paragraphs[0].add_run(str(lbl)); r.bold = True
            r.font.size = Pt(B["bs"] - 1); r.font.color.rgb = WHITE; r.font.name = B["hf"]
            shd(cl, PRI_HEX); cell_pad(cl, t=60, b=60)

    def cover_image_stream():
        if not openai_key:
            return None
        # IP-SAFE generic subject from enrichment (OpenAI refuses real people, celebrities,
        # brands, logos, copyrighted chars -> silent failure). Vary the ART STYLE by the
        # rotation seed so the cover is creative and different each run, not a fixed look.
        subj = (enrich.get("image_subject") or "").strip() or topic_clarified
        mood = {"PRIMARY": "playful, warm, friendly, full of wonder",
                "MIDDLE": "vibrant, dynamic and engaging",
                "SENIOR": "striking, sophisticated and refined"}[band_name]
        STYLES = ["bold flat-colour vector illustration", "lush storybook gouache painting",
                  "retro screen-print poster art", "papercut collage with layered depth",
                  "dynamic comic-book panel art", "clean isometric 3D-style render",
                  "textured chalk-and-pastel illustration", "vivid gradient digital art",
                  "hand-drawn ink-and-watercolour", "mid-century-modern poster design"]
        try:
            si = int(re.sub(r"\D", "", str(seed_raw)) or "0")
        except Exception:
            si = 0
        art = STYLES[si % len(STYLES)]
        prompt = ("A creative, eye-catching cover illustration of: " + subj + ". Art style: " + art +
                  ". Mood: " + mood + ". Be imaginative and visually striking -- strong composition, depth and "
                  "interesting lighting, never flat or generic. Show only an original, generic scene; do NOT depict "
                  "any real or named person, celebrity, athlete, brand, company, logo, trademark, or copyrighted "
                  "character. ABSOLUTELY NO TEXT, NO WORDS, NO LETTERS, NO NUMBERS anywhere. No borders, no watermark.")
        try:
            resp = requests.post(
                "https://api.openai.com/v1/images/generations",
                headers={"Authorization": "Bearer " + openai_key, "Content-Type": "application/json"},
                json={"model": "gpt-image-1-mini", "prompt": prompt, "size": "1024x1024", "n": 1},
                timeout=150)
            if resp.status_code != 200:
                return None
            b64 = resp.json()["data"][0]["b64_json"]
            return io.BytesIO(base64.b64decode(b64))
        except Exception:
            return None

    # ---------- COVER ----------
    pack_label = "Work Pack"
    meta = " / ".join(x for x in [yl, timeframe] if x)
    img = cover_image_stream()

    if band_name == "SENIOR":
        kicker("Literacy & Numeracy " + pack_label)
        pa(title, bold=True, sz=B["h0"], clr=PRI, after=4)
        p = doc.add_paragraph()
        p._p.get_or_add_pPr().append(parse_xml('<w:pBdr %s><w:bottom w:val="single" w:sz="10" w:space="2" w:color="%s"/></w:pBdr>' % (nsdecls("w"), PRI_HEX)))
        pa(topic_clarified, sz=B["bs"] + 3, clr=INK, before=6, after=2)
        pa(meta, it=True, clr=GY, after=16)
        if img:
            try:
                ip = doc.add_paragraph(); ip.alignment = WA.CENTER
                ip.add_run().add_picture(img, width=Cm(min(B["cov_w"], 10.5)))
            except Exception:
                pass
    else:
        for _ in range(2):
            doc.add_paragraph()
        pa(title, bold=True, sz=B["h0"], clr=PRI, al=WA.CENTER, after=2)
        pa(topic, sz=B["bs"] + 1, it=True, clr=SEC, al=WA.CENTER, after=10)
        if img:
            try:
                ip = doc.add_paragraph(); ip.alignment = WA.CENTER
                ip.add_run().add_picture(img, width=Cm(min(B["cov_w"], 11.0)))
            except Exception:
                pass
        pa(meta + "   |   " + pack_label, it=True, sz=B["bs"], clr=GY, al=WA.CENTER, after=14)

    # "belongs to" plate
    t = doc.add_table(rows=1, cols=1); cl = t.rows[0].cells[0]; cl.text = ""
    kp = cl.paragraphs[0]; chip(kp, "This work pack belongs to", PRI_HEX)
    np_ = cl.add_paragraph(); np_.paragraph_format.space_before = Pt(8)
    np_.add_run("Name: ").bold = True
    np_._p.get_or_add_pPr().append(parse_xml('<w:pBdr %s><w:bottom w:val="dotted" w:sz="6" w:space="6" w:color="B5B9BF"/></w:pBdr>' % nsdecls("w")))
    dp = cl.add_paragraph(); dp.paragraph_format.space_before = Pt(10)
    dp.add_run("Date: ").bold = True
    dp._p.get_or_add_pPr().append(parse_xml('<w:pBdr %s><w:bottom w:val="dotted" w:sz="6" w:space="6" w:color="B5B9BF"/></w:pBdr>' % nsdecls("w")))
    shd(cl, B["tint"] if B["box"] != "minimal" else "FFFFFF"); cell_borders(cl, left=PRI_HEX, box=GRID); cell_pad(cl)
    doc.add_page_break()

    # ---------- header / footer (body pages only) ----------
    sec0 = doc.sections[0]
    fp = sec0.footer.paragraphs[0]; fp.alignment = WA.CENTER
    fr = fp.add_run((topic + "   |   Page ")); fr.font.size = Pt(8); fr.font.color.rgb = GY
    fld = parse_xml('<w:fldSimple %s w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>' % nsdecls("w"))
    fp._p.append(fld)
    hp = sec0.header.paragraphs[0]; hp.alignment = WA.RIGHT
    hr = hp.add_run(yl + "  " + pack_label); hr.font.size = Pt(8); hr.font.color.rgb = GY

    # ---------- INTRO / welcome ----------
    if intro and not intro.startswith("ERR"):
        kicker("Welcome")
        addh("Before you begin", 1)
        for para in intro.split("\n"):
            if para.strip():
                pa(para.strip(), after=8)
        # warm-up mini-cards (fill the white space, ease into the activities).
        # Prefer enrichment.warmup; else compose from data we already have.
        warm = enrich.get("warmup") if isinstance(enrich.get("warmup"), list) else None
        if not warm:
            warm = []
            _facts = []
            for _d in (lit, num):
                if isinstance(_d, dict) and isinstance(_d.get("fun_facts"), list):
                    _facts += _d["fun_facts"]
            if _facts:
                warm.append({"kind": "fact", "text": str(_facts[0])})
            _vocab = enrich.get("key_vocabulary") if isinstance(enrich.get("key_vocabulary"), list) else []
            if _vocab:
                warm.append({"kind": "vocab", "text": "Words you'll meet: " + ", ".join(str(v) for v in _vocab[:4])})
            warm.append({"kind": "think", "text": "What do you already know about " + (topic or "this topic") + "? What would you like to find out?"})
        if isinstance(warm, list) and warm:
            doc.add_paragraph()
            kicker("Quick Warm-Up")
            kindcfg = {"fact": ("FFF8E7", "E8A33D", "DID YOU KNOW?"),
                       "vocab": ("FFFFFF", SEC_HEX, "WORDS TO KNOW"),
                       "think": (B["tint"], PRI_HEX, "QUICK THINK")}
            cards = [c for c in warm[:3] if isinstance(c, dict)]
            if cards:
                tw = doc.add_table(rows=1, cols=len(cards))
                tw.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, card in enumerate(cards):
                    fill, acc, lbl = kindcfg.get(card.get("kind", "fact"), kindcfg["fact"])
                    cl = tw.rows[0].cells[i]; cl.text = ""
                    p = cl.paragraphs[0]; chip(p, lbl, acc); p.add_run().add_break()
                    emit(cl.add_paragraph(), card.get("text", ""), sz=B["bs"] - 2)
                    if B["box"] != "minimal":
                        shd(cl, fill)
                    cell_borders(cl, left=acc, box=GRID); cell_pad(cl)
                cantsplit(tw)
        # No page break: let Section 1 flow straight on so the warm-up leads
        # into the activities and we don't leave half a blank page.
        doc.add_paragraph()

    # ---------- activity renderer ----------
    def render_glossary(act):
        g = act.get("glossary")
        if isinstance(g, list) and g:
            rows = ["**%s** -- %s" % (str(it.get("term", "")), str(it.get("definition", ""))) for it in g if isinstance(it, dict)]
            if rows:
                box(None, "glossary", rows=rows)

    def ra(act):
        if not isinstance(act, dict):
            return
        an = act.get("activity_number", 0)
        # H2 with a number chip
        h = doc.add_heading("", level=2)
        chip(h, str(an), PRI_HEX); h.add_run("  ")
        emit(h, act.get("activity_title", "Activity"), sz=None, clr=PRI)
        diff = act.get("difficulty", ""); mins = act.get("estimated_minutes", "")
        bits = [b for b in [diff, ("~%s min" % mins) if mins else ""] if b]
        if bits:
            pa("  ".join(bits), it=True, sz=B["bs"] - 3, clr=GY, after=2)
        if act.get("instructions"):
            ip = pa(act["instructions"], sz=B["bs"] - 1, after=4)
            ip.paragraph_format.keep_with_next = True
        render_glossary(act)
        c = act.get("content", {}) if isinstance(act.get("content"), dict) else {}
        at = act.get("activity_type", "")

        if c.get("passage"):
            box(c["passage"], "passage")
        if at == "cloze_passage":
            wb = c.get("word_bank", [])
            if isinstance(wb, list) and wb:
                box(" - ".join(str(w) for w in wb), "wordbank")
            ct = c.get("cloze_text", "")
            if ct:
                # render plain (no markdown) so stray * / blanks don't collide; normalise any blank style
                ct = str(ct).replace("**", "").replace("*", "")
                ct = re.sub(r"_{2,}(?:\d+_{2,})?|\b\d+\b(?=\s*_)", "_________", ct)
                ct = re.sub(r"_{3,}", "_________", ct)
                p = doc.add_paragraph(); rr = p.add_run(ct)
                rr.font.size = Pt(B["bs"]); rr.font.name = B["bf"]
        if at == "multiple_choice":
            lbl = ["A", "B", "C", "D"]
            for i, m in enumerate(c.get("mcq_items", []), 1):
                numq(i, m.get("stem", "") if isinstance(m, dict) else str(m))
                for j, opt in enumerate(m.get("options", []) if isinstance(m, dict) else []):
                    if j < 4:
                        op = doc.add_paragraph(); op.paragraph_format.left_indent = Cm(0.8); op.paragraph_format.space_after = Pt(2)
                        op.add_run("[  ]  %s.  " % lbl[j]).font.size = Pt(B["bs"] - 1)
                        emit(op, str(opt), sz=B["bs"] - 1)
        if at == "true_false":
            for i, tf in enumerate(c.get("tf_statements", []), 1):
                numq(i, tf.get("statement", "") if isinstance(tf, dict) else str(tf))
                tp = doc.add_paragraph(); tp.paragraph_format.left_indent = Cm(0.8)
                rr = tp.add_run("TRUE   /   FALSE"); rr.bold = True; rr.font.color.rgb = SEC
        if at == "matching":
            mp = c.get("match_pairs", {})
            if isinstance(mp, dict):
                left = mp.get("left", []); right = mp.get("right", []); mx = max(len(left), len(right), 1)
                t = doc.add_table(rows=mx + 1, cols=3); table_grid(t)
                hdr_cells(t.rows[0].cells, ["Item", "Your answer", "Choices"])
                abc = "ABCDEFGHIJKLMNOP"
                for ri in range(mx):
                    if ri < len(left):
                        cl = t.rows[ri + 1].cells[0]; cl.text = ""
                        emit(cl.paragraphs[0], "%d. %s" % (ri + 1, str(left[ri])), sz=B["bs"] - 1)
                    if ri < len(right):
                        cl = t.rows[ri + 1].cells[2]; cl.text = ""
                        emit(cl.paragraphs[0], "%s. %s" % (abc[ri], str(right[ri])), sz=B["bs"] - 1)
                    for cix in range(3):
                        cell_pad(t.rows[ri + 1].cells[cix], t=70, b=70)
        if at == "ordering":
            items = c.get("sequence_items", [])
            if isinstance(items, list) and items:
                box(" - ".join(str(x) for x in items), "plain")
                pa("Write the correct order:", bold=True, sz=B["bs"] - 1)
                for i in range(len(items)):
                    numq(i + 1, "______________________________________________")
        if at == "creative_writing":
            if c.get("writing_prompt"):
                box(c["writing_prompt"], "task")
            ws = c.get("writing_scaffold", [])
            if isinstance(ws, list):
                for si in ws:
                    doc.add_paragraph(str(si), style="List Bullet")
            wl = c.get("writing_lines", 8); wl = wl if isinstance(wl, int) else 8
            writelines(wl)
        if c.get("questions") and at not in ("cloze_passage", "multiple_choice", "true_false"):
            for i, q in enumerate(c["questions"], 1):
                numq(i, gq(q)); writelines(gl(q))
        if at == "calculation":
            calcs = c.get("calculations", [])
            if isinstance(calcs, list) and calcs:
                for i, cc in enumerate(calcs, 1):
                    numq(i, cc); writelines(1)
        if c.get("data_table") and at in ("data_interpretation", "drawing_graphing"):
            dt = c["data_table"]
            if isinstance(dt, dict):
                if dt.get("title"):
                    pa(dt["title"], bold=True, sz=B["bs"] - 1, clr=PRI)
                hdrs = dt.get("headers", []); rws = dt.get("rows", [])
                if hdrs and rws:
                    t = doc.add_table(rows=1 + len(rws), cols=len(hdrs)); table_grid(t)
                    hdr_cells(t.rows[0].cells, hdrs)
                    for ri, row in enumerate(rws):
                        for ci, val in enumerate(row):
                            if ci < len(hdrs):
                                cl2 = t.rows[ri + 1].cells[ci]; cl2.text = ""
                                emit(cl2.paragraphs[0], str(val), sz=B["bs"] - 1); cell_pad(cl2, t=60, b=60)
                        if ri % 2 == 1:
                            for ci in range(len(hdrs)):
                                shd(t.rows[ri + 1].cells[ci], B["tint"])
        if at == "drawing_graphing":
            if c.get("graph_instructions"):
                pa(c["graph_instructions"], bold=True, sz=B["bs"] - 1)
            gh = c.get("graph_height_cm", 6); gh = gh if isinstance(gh, int) else 6
            t = doc.add_table(rows=1, cols=1); table_grid(t)
            t.rows[0].height = Cm(gh); t.rows[0].cells[0].text = ""
        if at == "questions" and c.get("questions"):
            for i, q in enumerate(c["questions"], 1):
                numq(i, gq(q)); writelines(gl(q))
        if act.get("challenge_box"):
            box(act["challenge_box"], "challenge")

    # ---------- section renderer ----------
    def divider(num_label, title):
        for _ in range(4):
            doc.add_paragraph()
        kp = doc.add_paragraph(); kp.alignment = WA.CENTER
        r = kp.add_run(num_label.upper()); r.bold = True; r.font.size = Pt(B["bs"] - 2); r.font.color.rgb = SEC
        r.font.name = B["hf"]; r._r.get_or_add_rPr().append(parse_xml('<w:spacing %s w:val="60"/>' % nsdecls("w")))
        tp = doc.add_paragraph(); tp.alignment = WA.CENTER
        emit(tp, title, sz=B["h0"] - 4, clr=PRI)
        for rn in tp.runs:
            rn.bold = True
        doc.add_page_break()

    def rs(data, sec_no):
        if not _ok(data):
            addh("Section %d" % sec_no, 1)
            pa("This section could not be generated -- please run the tool again.", clr=PRI, bold=True)
            doc.add_page_break(); return
        title = data.get("section_title", "Section %d" % sec_no)
        # nice section header: kicker + heading + rule
        kicker("Section %d" % sec_no)
        addh(title.split(" - ")[-1].strip("'\" ") or title, 1)
        p = doc.add_paragraph()
        p._p.get_or_add_pPr().append(parse_xml('<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="2" w:color="%s"/></w:pBdr>' % (nsdecls("w"), PRI_HEX)))
        if data.get("fun_intro"):
            box(data["fun_intro"], "plain")
        facts = data.get("fun_facts", []); facts = facts if isinstance(facts, list) else []
        acts = data.get("activities", []); acts = acts if isinstance(acts, list) else []
        days = {}
        for a in acts:
            if isinstance(a, dict):
                d = a.get("day", 1); d = d if (isinstance(d, int) and d >= 1) else 1
                days.setdefault(d, []).append(a)
        multi = is_multi and len(days) > 1; fi = 0
        for dn in sorted(days.keys()):
            if multi:
                addh("Day %d" % dn, 3)
            for act in days[dn]:
                ra(act)
                an = act.get("activity_number", 0)
                if isinstance(an, int) and fi < len(facts) and an % 2 == 0:
                    box(str(facts[fi]), "fact"); fi += 1
            if multi:
                doc.add_page_break()
        for f in facts[fi:]:
            box(str(f), "fact")
        if data.get("reflection_prompt"):
            box(data["reflection_prompt"], "reflection"); writelines(3)
        doc.add_page_break()

    # ---------- build body ----------
    rs(lit, 1)
    if _ok(num):
        divider("Section 2", (num.get("section_title", "Numeracy").split(" - ")[-1].strip("'\" ")) if isinstance(num, dict) else "Numeracy")
        rs(num, 2)
    else:
        rs(num, 2)

    # ---------- teacher answer key ----------
    if any(isinstance(d, dict) and d.get("answer_key") for d in (lit, num)):
        divider("Teacher Section", "Answer Key")
        pa("For teacher use only -- remove before printing the student copy.", it=True, sz=B["bs"] - 2, clr=GY)
        for d in (lit, num):
            if not isinstance(d, dict) or not d.get("answer_key"):
                continue
            pa((d.get("section_title", "")).split(" - ")[-1].strip("'\" "), bold=True, clr=PRI, sz=B["bs"])
            ak = d["answer_key"]
            if not isinstance(ak, list):
                continue
            for a in ak:
                if not isinstance(a, dict):
                    continue
                rows = ["**Activity %s: %s**" % (str(a.get("activity_number", "")), a.get("activity_title", ""))]
                ans = a.get("answers", [])
                ans = ans if isinstance(ans, list) else [str(ans)]
                rows += [str(x) for x in ans]
                box(None, "plain", rows=rows)

    pa("Generated by Teachers Assistant AI", it=True, sz=8, clr=GY, al=WA.RIGHT, before=12)
    doc.save("workpack.docx")
    RESULT = insert_temp_file("workpack.docx")
